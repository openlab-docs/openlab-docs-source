# -*- coding: utf-8 -*-
"""
    sphinxcontrib-docxwriter
    ~~~~~~~~~~~~~~~~~~~~~~~~~~

    Custom docutils writer for OpenXML (docx).

    :copyright:
        Copyright 2010 by shimizukawa at gmail dot com (Sphinx-users.jp).
    :license: BSD, see LICENSE for details.
"""
from __future__ import division
import os
import sys

# noinspection PyUnresolvedReferences
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
# noinspection PyProtectedMember
from docx.table import _Cell
from docx import Document

from docutils import nodes, writers

try:
    from PIL import Image
except ImportError as exp:
    Image = None

import logging
logging.basicConfig(
    filename='docx.log',
    filemode='w',
    level=logging.INFO,
    format="%(asctime)-15s  %(message)s"
)
logger = logging.getLogger('docx')


def dprint(_func=None, **kw):
    """Print debug information."""
    logger.info('-'*50)
    # noinspection PyProtectedMember
    f = sys._getframe(1)
    if kw:
        text = ', '.join('%s = %s' % (k, v) for k, v in kw.items())
    else:
        text = dict((k, repr(v)) for k, v in f.f_locals.items() if k != 'self')
        text = str(text)

    if _func is None:
        _func = f.f_code.co_name

    print(text)
    logger.info(' '.join([_func, text]))


# noinspection PyUnusedLocal
def _make_depart_admonition(name):
    # noinspection PyMissingOrEmptyDocstring,PyUnusedLocal
    def depart_admonition(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
    return depart_admonition


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring
class DocxWriter(writers.Writer):
    """docutil writer class for docx files"""
    supported = ('docx',)
    settings_spec = ('No options here.', '', ())
    settings_defaults = {}

    output = None
    template_dir = "NO"

    def __init__(self, builder):
        writers.Writer.__init__(self)
        self.builder = builder
        self.template_setup()  # setup before call almost docx methods.

        if self.template_dir == "NO":
            dc = Document()
        else:
            dc = Document(os.path.join('source', self.template_dir))
        self.docx_container = dc

    def template_setup(self):
        dotx = self.builder.config['docx_template']
        if dotx:
            logger.info("MK using template {}".format(dotx))
            self.template_dir = dotx

    def save(self, filename):
        self.docx_container.save(filename)

    def translate(self):
        visitor = DocxTranslator(
                self.document, self.builder, self.docx_container)
        self.document.walkabout(visitor)
        self.output = ''  # visitor.body


class DocxState(object):
    """
    DocxState class keeps track of which part of the document is being worked on.

    In particular it is used to allow lists in tables.
    """
    def __init__(self, location=None):
        self.location = location
        self.table = None
        self.column_widths = None
        self.table_style = None
        self.more_cols = 0
        self.row = None
        self.cell_counter = 0
        self.ncolumns = 1
        "Number of columns in the current table."


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring,PyUnusedLocal
class DocxTranslator(nodes.NodeVisitor):
    """Visitor class to create docx content."""

    def __init__(self, document, builder, docx_container):
        self.builder = builder
        self.docx_container = docx_container
        nodes.NodeVisitor.__init__(self, document)

        # TODO: Perhaps move the list_style into DocxState.
        # However, it should still be a list, and not a separate state,
        # because nested lists are not really nested.
        # So it will only be necessary if there are lists in tables
        # that are in lists.
        self.list_style = []
        self.list_level = 0

        # TODO: And what about sectionlevel?
        self.sectionlevel = 0

        self.table_style_default = 'Grid Table 4 Accent 1'
        self.in_literal_block = False
        self.strong = False
        self.emphasis = False

        self.current_state = DocxState(location=self.docx_container)
        self.current_state.table_style = self.table_style_default

        "The place where paragraphs will be added."
        self.old_states = []
        "A list of older states, e.g. typically [document, table-cell]"

        self.current_paragraph = None
        "The current paragraph that text is being added to."

    def add_text(self, text):
        dprint()
        textrun = self.current_paragraph.add_run(text)
        if self.strong:
            textrun.bold = True
        if self.emphasis:
            textrun.italic = True

    def new_state(self, location):
        dprint()
        self.old_states.append(self.current_state)
        self.current_state = DocxState(location=location)
        # print("HB new_state: {}".format(len(self.old_states)))

    def end_state(self, first=None):
        dprint()
        self.current_state = self.old_states.pop()
        # print("HB old_state: {}".format(len(self.old_states)))

    def visit_start_of_file(self, node):
        dprint()
        # TODO: HB should visit_start_of_file reset the sectionlevel?
        # If so, should it start a new state? If so, with which location?

        # FIXME: visit_start_of_file not close previous section.
        # sectionlevel keep previous and new file's heading level start with
        # previous + 1.
        # This quick hack reset sectionlevel per file.
        # (BTW Sphinx has heading levels per file? or entire document?)
        self.sectionlevel = 0

    def depart_start_of_file(self, node):
        dprint()

    def visit_document(self, node):
        print('visit_document')
        dprint()

    def depart_document(self, node):
        dprint()

    def visit_highlightlang(self, node):
        dprint()
        print('visit_highlightlang')
        raise nodes.SkipNode

    def visit_section(self, node):
        print('visit_section')
        dprint()
        self.sectionlevel += 1

    def depart_section(self, node):
        dprint()
        if self.sectionlevel > 0:
            self.sectionlevel -= 1

    def visit_topic(self, node):
        dprint()
        print('visit_topic')
        raise nodes.SkipNode

    def depart_topic(self, node):
        dprint()
        raise nodes.SkipNode

    visit_sidebar = visit_topic
    depart_sidebar = depart_topic

    def visit_rubric(self, node):
        dprint()
        print('visit_rubric')
        raise nodes.SkipNode
        # self.add_text('-[ ')

    def depart_rubric(self, node):
        dprint()
        print('depart_rubric')
        raise nodes.SkipNode
        # self.add_text(' ]-')

    def visit_compound(self, node):
        dprint()
        pass

    def depart_compound(self, node):
        dprint()
        pass

    def visit_glossary(self, node):
        dprint()
        pass

    def depart_glossary(self, node):
        dprint()
        pass

    def visit_title(self, node):
        for i in self.docx_container.styles:
            print('{} || {}'.format(i.name, i))
        
        print('visit_title')
        dprint()
        self.current_paragraph = self.current_state.location.add_heading(level=self.sectionlevel)

    def depart_title(self, node):
        dprint()

    def visit_subtitle(self, node):
        print('visit_subtitle')
        dprint()
        pass

    def depart_subtitle(self, node):
        dprint()
        pass

    def visit_attribution(self, node):
        dprint()
        print('visit_attribution')
        raise nodes.SkipNode
        # self.add_text('-- ')

    def depart_attribution(self, node):
        dprint()
        pass

    def visit_desc(self, node):
        dprint()
        pass

    def depart_desc(self, node):
        dprint()
        pass

    def visit_desc_signature(self, node):
        dprint()
        print('visit_desc_signature')
        raise nodes.SkipNode

    def depart_desc_signature(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_desc_name(self, node):
        dprint()
        pass

    def depart_desc_name(self, node):
        dprint()
        pass

    def visit_desc_addname(self, node):
        dprint()
        pass

    def depart_desc_addname(self, node):
        dprint()
        pass

    def visit_desc_type(self, node):
        dprint()
        pass

    def depart_desc_type(self, node):
        dprint()
        pass

    def visit_desc_returns(self, node):
        dprint()
        print('visit_desc_returns')
        raise nodes.SkipNode
        # self.add_text(' -> ')

    def depart_desc_returns(self, node):
        dprint()
        pass

    def visit_desc_parameterlist(self, node):
        dprint()
        print('visit_desc_parameterlist')
        raise nodes.SkipNode
        # self.add_text('(')
        # self.first_param = 1

    def depart_desc_parameterlist(self, node):
        dprint()
        print('depart_desc_parameterlist')
        raise nodes.SkipNode
        # self.add_text(')')

    def visit_desc_parameter(self, node):
        dprint()
        print('visit_desc_parameter')
        raise nodes.SkipNode
        # if not self.first_param:
        #     self.add_text(', ')
        # else:
        #     self.first_param = 0
        # self.add_text(node.astext())
        # raise nodes.SkipNode

    def visit_desc_optional(self, node):
        dprint()
        print('visit_desc_optional')
        raise nodes.SkipNode
        # self.add_text('[')

    def depart_desc_optional(self, node):
        dprint()
        print('depart_desc_optional')
        raise nodes.SkipNode
        # self.add_text(']')

    def visit_desc_annotation(self, node):
        dprint()
        pass

    def depart_desc_annotation(self, node):
        dprint()
        pass

    def visit_refcount(self, node):
        dprint()
        pass

    def depart_refcount(self, node):
        dprint()
        pass

    def visit_desc_content(self, node):
        dprint()
        print('visit_desc_content')
        raise nodes.SkipNode
        # self.add_text('\n')

    def depart_desc_content(self, node):
        dprint()
        print('depart_desc_content')
        raise nodes.SkipNode

    def visit_figure(self, node):
        # FIXME: figure text become normal paragraph instead of caption.
        dprint()

    def depart_figure(self, node):
        dprint()

    def visit_caption(self, node):
        dprint()
        pass

    def depart_caption(self, node):
        dprint()
        pass

    def visit_productionlist(self, node):
        dprint()
        print('visit_productionlist')
        raise nodes.SkipNode
        # names = []
        # for production in node:
        #     names.append(production['tokenname'])
        # maxlen = max(len(name) for name in names)
        # for production in node:
        #     if production['tokenname']:
        #         self.add_text(production['tokenname'].ljust(maxlen) + ' ::=')
        #         lastname = production['tokenname']
        #     else:
        #         self.add_text('%s    ' % (' '*len(lastname)))
        #     self.add_text(production.astext() + '\n')
        # raise nodes.SkipNode

    def visit_seealso(self, node):
        dprint()

    def depart_seealso(self, node):
        dprint()

    def visit_footnote(self, node):
        dprint()
        print('visit_footnote')
        raise nodes.SkipNode
        # self._footnote = node.children[0].astext().strip()

    def depart_footnote(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_citation(self, node):
        dprint()
        print('visit_citation')
        raise nodes.SkipNode
        # if len(node) and isinstance(node[0], nodes.label):
        #     self._citlabel = node[0].astext()
        # else:
        #     self._citlabel = ''

    def depart_citation(self, node):
        dprint()
        print('depart_citation')
        raise nodes.SkipNode

    def visit_label(self, node):
        dprint()
        print('visit_label')
        raise nodes.SkipNode

    # XXX: option list could use some better styling

    def visit_option_list(self, node):
        dprint()
        pass

    def depart_option_list(self, node):
        dprint()
        pass

    def visit_option_list_item(self, node):
        dprint()
        print('visit_option_list_item')
        raise nodes.SkipNode

    def depart_option_list_item(self, node):
        dprint()
        print('depart_option_list_item')
        raise nodes.SkipNode

    def visit_option_group(self, node):
        dprint()
        print('visit_option_group')
        raise nodes.SkipNode
        # self._firstoption = True

    def depart_option_group(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('     ')

    def visit_option(self, node):
        dprint()
        print('visit_option')
        raise nodes.SkipNode
        # if self._firstoption:
        #     self._firstoption = False
        # else:
        #     self.add_text(', ')

    def depart_option(self, node):
        dprint()
        pass

    def visit_option_string(self, node):
        dprint()
        pass

    def depart_option_string(self, node):
        dprint()
        pass

    def visit_option_argument(self, node):
        dprint()
        print('visit_option_argument')
        raise nodes.SkipNode
        # self.add_text(node['delimiter'])

    def depart_option_argument(self, node):
        dprint()
        pass

    def visit_description(self, node):
        dprint()
        pass

    def depart_description(self, node):
        dprint()
        pass

    def visit_tabular_col_spec(self, node):
        dprint()
        # TODO: properly implement this!!
        spec = node['spec']
        widths = [float(l.split('cm')[0]) for l in spec.split("{")[1:]]
        self.current_state.column_widths = widths
        print('visit_tabular_col_spec')
        raise nodes.SkipNode

    def visit_colspec(self, node):
        dprint()
        # The difficulty here is getting the right column width.
        # This can be specified with a tabular_col_spec, see above.
        #
        # Otherwise it is derived from the number of columns, which is
        # defined in visit_tgroup (a bit hackish).
        # The _block_width is the full width of the document, and this
        # is divided by the number of columns.
        #
        # It would perhaps also be possible to use node['colwidth'] in some way.
        # node['colwidth'] contains an integer like 22, the width of the column in ascii
        if self.current_state.column_widths:
            width = self.current_state.column_widths[0]
            self.current_state.column_widths = self.current_state.column_widths[1:]
            col = self.current_state.table.add_column(Cm(width))
        else:
            # noinspection PyProtectedMember
            col = self.current_state.table.add_column(self.docx_container._block_width // self.current_state.ncolumns)

        print('visit_colspec')
        raise nodes.SkipNode

    def depart_colspec(self, node):
        dprint()

    def visit_tgroup(self, node):
        dprint()
        colspecs = [c for c in node.children if isinstance(c, nodes.colspec)]
        self.current_state.ncolumns = len(colspecs)
        # print("HB VT {} {} {}".format(
        #     self.current_state.ncolumns, len(node.children), [type(c) for c in node.children]))

    def depart_tgroup(self, node):
        dprint()
        self.current_state.ncolumns = 1
        pass

    def visit_thead(self, node):
        dprint()
        print('visit_thead')
        pass

    def depart_thead(self, node):
        dprint()
        pass

    def visit_tbody(self, node):
        dprint()

    def depart_tbody(self, node):
        dprint()
        pass

    def visit_row(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()
        self.current_state.cell_counter = 0

    def depart_row(self, node):
        dprint()
        pass

    def visit_entry(self, node):
        dprint()
        if 'morerows' in node:
            raise NotImplementedError('Row spanning cells are not implemented.')
        if 'morecols' in node:
            # Hack to make column spanning possible. TODO FIX
            self.current_state.more_cols = node['morecols']

        cell = self.current_state.row.cells[self.current_state.cell_counter]
        # A new paragraph will be added by Sphinx, so remove the automated one
        # This turns out to be not possible, so instead the existing one is
        # reused in visit_paragraph.
        # cell.paragraphs.pop()
        if self.current_state.more_cols:
            # Perhaps this commented line works no too.
            # cell = cell.merge(self.row.cells[self.cell_counter + self.more_cols])
            for i in range(self.current_state.more_cols):
                cell = cell.merge(self.current_state.row.cells[self.current_state.cell_counter + i + 1])

        self.new_state(location=cell)
        # For some annoying reason, a new paragraph is automatically added
        # to each table cell. This is frustrating when you want, e.g. to
        # add a list item instead of a normal paragraph.
        self.current_paragraph = cell.paragraphs[0]

    def depart_entry(self, node):
        dprint()
        self.end_state()
        self.current_state.cell_counter = self.current_state.cell_counter + self.current_state.more_cols + 1
        self.current_state.more_cols = 0

    def visit_table(self, node):
        dprint()

        style = self.current_state.table_style
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.TABLE)
        except KeyError as exc:
            print(exc)
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        # Columns are added when a colspec is visited.

        # It is only possible to use a style in add_table when adding a
        # table to the root document. That is, not for a table in a table.
        if len(self.old_states):
            self.current_state.table = self.current_state.location.add_table(rows=0, cols=0)
        else:
            self.current_state.table = self.current_state.location.add_table(
                rows=0, cols=0, style=style)

    def depart_table(self, node):
        dprint()

        self.current_state.table = None
        self.current_state.table_style = self.table_style_default

        # Add an empty paragraph to prevent tables from being concatenated.
        # TODO: Figure out some better solution.
        self.current_state.location.add_paragraph("")

    def visit_acks(self, node):
        dprint()
        print('visit_acks')
        raise nodes.SkipNode
        # self.add_text(', '.join(n.astext() for n in node.children[0].children)
        #               + '.')

    def visit_image(self, node):
        dprint()
        uri = node.attributes['uri']
        file_path = os.path.join(self.builder.env.srcdir, uri)
        width, height = self.get_image_scaled_width_height(node, file_path)

        from docx.shared import Pt as Cm
        self.docx_container.add_picture(file_path, width=Cm(width), height=Cm(height))

    def get_image_scale(self, node):
        if 'scale' in node.attributes:
            try:
                scale = int(node.attributes['scale'])
                if scale < 1: # or scale > 100:
                    self.document.reporter.warning(
                        'scale out of range (%s), using 1.' % (scale, ))
                    scale = 1
                scale = scale * 0.01
            except ValueError as e:
                self.document.reporter.warning(
                    'Invalid scale for image: "%s"' % (
                        node.attributes['scale'], ))
        else:
            scale = 1.0
        return scale

    def get_image_scaled_width_height(self, node, filename):
        dpi = (72, 72)

        if Image is not None :
            try:
              imageobj = Image.open(filename, 'r')
            except:
              raise RuntimeError('Fail to open image file: %s' % filename)

            dpi = imageobj.info.get('dpi', dpi)
            # dpi information can be (xdpi, ydpi) or xydpi
            try: iter(dpi)
            except: dpi = (dpi, dpi)
        else:
            imageobj = None
            raise RuntimeError('image size not fully specified and PIL not installed')

        scale = self.get_image_scale(node)
        width = self.get_image_width_height(node, 'width')
        height = self.get_image_width_height(node, 'height')

        if width is not None and width[1] == '%':
           width = [int(self.docx.styleDocx.document_width * width[0] * 0.00284 ), 'px']

        if height is not None and height[1] == '%':
           height = [int(self.docx.styleDocx.document_height * height[0] * 0.00284 ), 'px']

        if width is None or height is None:
            if imageobj is None:
                raise RuntimeError(
                    'image size not fully specified and PIL not installed')
            if width is None:
                if height is None:
                     width = [imageobj.size[0], 'px']
                     height = [imageobj.size[1], 'px']
                else:
                     scaled_width = imageobj.size[0] * height[0] /imageobj.size[1]
                     width = [scaled_width, 'px']
            else:
                if height is None:
                     scaled_height = imageobj.size[1] * width[0] / imageobj.size[0]
                     height = [scaled_height, 'px']
                else:
                     height = [imageobj.size[1], 'px']

        width[0] *= scale
        height[0] *= scale
        if width[1] == 'in': width = [width[0] * dip[0], 'px']
        if height[1] == 'in': height = [height[0] *dip[1], 'px']

        #  We shoule shulink image (multiply 72/96)
        width[0] *= 0.75
        height[0] *= 0.75

        #
        # maxwidth = int(self.docx_container.sections[0].page_width  * 0.284 ) * 0.9
        maxwidth = 456

        if width[0] > maxwidth :
            ratio = height[0] / width[0]
            width[0] = maxwidth
            height[0] = width[0] * ratio

        maxheight = int(self.docx_container.sections[0].page_height  * 0.284 ) * 0.9
        if height[0] > maxheight :
            ratio = width[0] / height[0]
            height[0] = maxheight
            width[0] = height[0] * ratio

        return int(width[0]), int(height[0])

    def get_image_width_height(self, node, attr):
        size = None
        if attr in node.attributes:
          size = node.attributes[attr]
          if size[-1] == '%' :
            size = float(size[:-1])
            size = [size, '%']
          else:
            unit = size[-2:]
            if unit.isalpha():
                size = size[:-2]
            else:
                unit = 'px'
            try:
                size = float(size)
            except ValueError as e:
                self.document.reporter.warning(
                    'Invalid %s for image: "%s"' % (
                        attr, node.attributes[attr]))
            size = [size, unit]
        return size


    def depart_image(self, node):
        dprint()

    def visit_transition(self, node):
        dprint()
        print('visit_transition')
        raise nodes.SkipNode
        # self.add_text('=' * 70)

    def visit_bullet_list(self, node):
        dprint()
        # TODO: Apparently it is necessary to take into account whether
        # the list is numbered or not, like the original code did.
        # But that code did not properly account for the level.
        # So merge these two attempts.
        # self.list_style.append('ListBullet')
        self.list_level += 1

    def depart_bullet_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.append('ListNumber')
        self.list_level += 1

    def depart_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_definition_list(self, node):
        dprint()
        print('visit_definition_list')
        raise nodes.SkipNode
        # self.list_style.append(-2)

    def depart_definition_list(self, node):
        dprint()
        raise nodes.SkipNode
        # self.list_style.pop()

    def visit_list_item(self, node):
        dprint()
        # A new paragraph is created here, but the next visit is to
        # paragraph, so that would add another paragraph. That is
        # prevented if current_paragraph is an empty List paragraph.
        style = 'List Bullet' if self.list_level < 2 else 'List Bullet {}'.format(self.list_level)
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        curloc = self.current_state.location
        if isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    # print("HB VLI reuse {}".format(style))
                    self.current_paragraph = curloc.paragraphs[0]
                    self.current_paragraph.style = style
                else:
                    # print("HB VLI create 1 {}".format(style))
                    self.current_paragraph = curloc.add_paragraph(style=style)
            else:
                # print("HB VLI create 2 {}".format(style))
                self.current_paragraph = curloc.add_paragraph(style=style)
        else:
            # print("HB VLI create 3 {}".format(style))
            self.current_paragraph = curloc.add_paragraph(style=style)

        # print("HB VLI end {}".format(self.current_paragraph.style))

    def depart_list_item(self, node):
        dprint()

    def visit_definition_list_item(self, node):
        dprint()
        print('visit_definition_list_item')
        raise nodes.SkipNode

    def depart_definition_list_item(self, node):
        dprint()
        pass

    def visit_term(self, node):
        dprint()
        print('visit_term')
        raise nodes.SkipNode

    def depart_term(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_classifier(self, node):
        dprint()
        print('visit_classifier')
        raise nodes.SkipNode
        # self.add_text(' : ')

    def depart_classifier(self, node):
        dprint()
        print('depart_classifier')
        raise nodes.SkipNode

    def visit_definition(self, node):
        dprint()
        print('visit_definition')
        raise nodes.SkipNode

    def depart_definition(self, node):
        dprint()
        print('depart_definition')
        raise nodes.SkipNode

    def visit_field_list(self, node):
        dprint()
        pass

    def depart_field_list(self, node):
        dprint()
        pass

    def visit_field(self, node):
        dprint()
        pass

    def depart_field(self, node):
        dprint()
        pass

    def visit_field_name(self, node):
        dprint()
        print('visit_field_name')
        raise nodes.SkipNode

    def depart_field_name(self, node):
        dprint()
        print('depart_field_name')
        raise nodes.SkipNode
        # self.add_text(':')

    def visit_field_body(self, node):
        dprint()
        print('visit_field_body')
        raise nodes.SkipNode

    def depart_field_body(self, node):
        dprint()
        print('depart_field_body')
        raise nodes.SkipNode

    def visit_centered(self, node):
        dprint()
        pass

    def depart_centered(self, node):
        dprint()
        pass

    def visit_hlist(self, node):
        dprint()
        pass

    def depart_hlist(self, node):
        dprint()
        pass

    def visit_hlistcol(self, node):
        dprint()
        pass

    def depart_hlistcol(self, node):
        dprint()
        pass

    def visit_admonition(self, node):
        dprint()
        print('visit_admonition')
        raise nodes.SkipNode

    def depart_admonition(self, node):
        dprint()
        print('depart_admonition')
        raise nodes.SkipNode

    def _visit_admonition(self, node):
        dprint()
        print('_visit_admonition')
        raise nodes.SkipNode

    visit_attention = _visit_admonition
    depart_attention = _make_depart_admonition('attention')
    visit_caution = _visit_admonition
    depart_caution = _make_depart_admonition('caution')
    visit_danger = _visit_admonition
    depart_danger = _make_depart_admonition('danger')
    visit_error = _visit_admonition
    depart_error = _make_depart_admonition('error')
    visit_hint = _visit_admonition
    depart_hint = _make_depart_admonition('hint')
    visit_important = _visit_admonition
    depart_important = _make_depart_admonition('important')
    visit_note = _visit_admonition
    depart_note = _make_depart_admonition('note')
    visit_tip = _visit_admonition
    depart_tip = _make_depart_admonition('tip')
    visit_warning = _visit_admonition
    depart_warning = _make_depart_admonition('warning')

    def visit_versionmodified(self, node):
        dprint()
        print('visit_versionmodified')
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
        # if node.children:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + ': ')
        # else:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + '.')

    def depart_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_literal_block(self, node):
        dprint()
        # TODO: Check whether literal blocks work in tables and lists.
        self.in_literal_block = True

        # Unlike with Lists, there will not be a visit to paragraph in a
        # literal block, so we *must* create the paragraph here.
        style = 'code'

        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        self.current_paragraph = self.current_state.location.add_paragraph(style=style)
        self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def depart_literal_block(self, node):
        dprint()
        self.in_literal_block = False

    def visit_doctest_block(self, node):
        dprint()
        print('visit_doctest_block')
        raise nodes.SkipNode

    def depart_doctest_block(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_line_block(self, node):
        dprint()
        print('visit_line_block')
        # raise nodes.SkipNode

    def depart_line_block(self, node):
        dprint()
        # raise nodes.SkipNode

    def visit_line(self, node):
        dprint()
        print('visit_line')
        style = 'visit_line'

        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        self.current_paragraph = self.current_state.location.add_paragraph(style=style)
        self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def depart_line(self, node):
        dprint()
        pass

    def visit_block_quote(self, node):
        dprint()

    def depart_block_quote(self, node):
        dprint()

    def visit_compact_paragraph(self, node):
        dprint()

    def depart_compact_paragraph(self, node):
        dprint()

    def visit_paragraph(self, node):
        dprint()

        curloc = self.current_state.location

        if 'List' in self.current_paragraph.style.name and not self.current_paragraph.text:
            # print("HB VP List")
            # This is the first paragraph in a list item, so do not create another one.
            pass
        elif isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                    # print("HB VP cell reuse")
                else:
                    self.current_paragraph = curloc.add_paragraph()
                    # print("HB VP cell create 1")
            else:
                self.current_paragraph = curloc.add_paragraph()
                # print("HB VP cell create 2")
            # HACK because the style is messed up, TODO FIX
            self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.current_paragraph.paragraph_format.left_indent = 0
        else:
            # print("HB VP normal")
            self.current_paragraph = curloc.add_paragraph()

        # print("HB VP end {}".format(self.current_paragraph.style))

    def depart_paragraph(self, node):
        dprint()

    def visit_target(self, node):
        dprint()
        print('visit_target')
        raise nodes.SkipNode

    def visit_index(self, node):
        dprint()
        print('visit_index')
        raise nodes.SkipNode

    def visit_substitution_definition(self, node):
        dprint()
        print('visit_substitution_definition')
        raise nodes.SkipNode

    def visit_pending_xref(self, node):
        dprint()
        pass

    def depart_pending_xref(self, node):
        dprint()
        pass

    def visit_reference(self, node):
        print('visit_reference')
        dprint()
        pass

    def depart_reference(self, node):
        dprint()
        pass

    def visit_download_reference(self, node):
        dprint()
        pass

    def depart_download_reference(self, node):
        dprint()
        pass

    def visit_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = True

    def depart_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = False

    def visit_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def depart_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def visit_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = True

    def depart_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = False

    def visit_abbreviation(self, node):
        dprint()
        # self.add_text('')

    def depart_abbreviation(self, node):
        dprint()
        # if node.hasattr('explanation'):
        #     self.add_text(' (%s)' % node['explanation'])

    def visit_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def depart_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def visit_literal(self, node):
        dprint()
        # self.add_text('``')

    def depart_literal(self, node):
        dprint()
        # self.add_text('``')

    def visit_subscript(self, node):
        dprint()
        print('visit_subscript')
        raise nodes.SkipNode
        # self.add_text('_')

    def depart_subscript(self, node):
        dprint()
        pass

    def visit_superscript(self, node):
        dprint()
        print('visit_superscript')
        raise nodes.SkipNode
        # self.add_text('^')

    def depart_superscript(self, node):
        dprint()
        pass

    def visit_footnote_reference(self, node):
        dprint()
        print('visit_footnote_reference')
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_citation_reference(self, node):
        dprint()
        print('visit_citation_reference')
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_Text(self, node):
        dprint()
        text = node.astext()
        if not self.in_literal_block:
            # assert '\n\n' not in text, 'Found \n\n'
            # Replace double enter with single enter, and single enter with space.
            string_magic = 'TWOENTERSMAGICSTRING'
            text = text.replace('\n\n', string_magic)
            text = text.replace('\n', ' ')
            text = text.replace(string_magic, '\n')
        self.add_text(text)

    def depart_Text(self, node):
        dprint()
        pass

    def visit_generated(self, node):
        dprint()
        pass

    def depart_generated(self, node):
        dprint()
        pass

    def visit_inline(self, node):
        dprint()
        pass

    def depart_inline(self, node):
        dprint()
        pass

    def visit_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('>>')

    def depart_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<<')

    def visit_system_message(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<SYSTEM MESSAGE: %s>' % node.astext())

    def visit_comment(self, node):
        dprint()
        # TODO: FIX Dirty hack / kludge to set table style.
        # Use proper directives or something like that
        comment = node[0]
        if 'DocxTableStyle' in comment:
            self.current_state.table_style = comment.split('DocxTableStyle')[-1].strip()
        # print("HB tablestyle {}".format(self.current_state.table_style))
        raise nodes.SkipNode

    def visit_meta(self, node):
        dprint()
        raise nodes.SkipNode
        # only valid for HTML

    def visit_raw(self, node):
        dprint()
        raise nodes.SkipNode
        # if 'text' in node.get('format', '').split():
        #     self.body.append(node.astext())

    def unknown_visit(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)

    def unknown_departure(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)
